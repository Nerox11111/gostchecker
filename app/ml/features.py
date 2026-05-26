import re
from pathlib import Path

from docx import Document
from docx.shared import Mm


FEATURES = [
    "page_count",
    "heading_depth",
    "has_toc",
    "has_title",
    "has_abstract",
    "has_intro",
    "has_conclusion",
    "table_count",
    "figure_count",
    "formula_count",
    "ref_count",
    "appendix_count",
    "correct_margins",
    "char_count",
    "avg_para_len",
    "keyword_score",
]


CLASS_KEYWORDS = {
    "lab_work": {"лабораторная", "лабораторной", "опыт", "измерение"},
    "practice": {"практическая", "практикум", "задание"},
    "coursework": {"курсовая", "курсовой", "проект"},
    "internship": {"практика", "дневник", "отчет по практике"},
    "thesis_bachelor": {"бакалавр", "выпускная квалификационная", "вкр"},
    "thesis_master": {"магистр", "магистерская", "диссертация"},
    "scientific_rpt": {"научно-исследовательская", "исследование", "реферат"},
    "rnd_nir": {"нир", "научно-исследовательская работа", "отчет о нир"},
}


def _paragraph_texts(document: Document) -> list[str]:
    return [p.text.strip() for p in document.paragraphs if p.text.strip()]


def _heading_depth(text: str) -> int:
    match = re.match(r"^\s*(\d+(?:\.\d+){0,3})(?!\.)\s+\S+", text)
    if not match:
        return 0
    return match.group(1).count(".") + 1


def _ref_count(texts: list[str]) -> int:
    start = None
    for idx, text in enumerate(texts):
        upper = text.upper()
        if "СПИСОК" in upper and ("ИСТОЧНИ" in upper or "ЛИТЕРАТУР" in upper):
            start = idx + 1
            break
    if start is None:
        return 0
    count = 0
    for text in texts[start:]:
        if re.match(r"^\s*(\d+[\).\s]|[-–—])", text):
            count += 1
        elif text.upper().startswith("ПРИЛОЖЕНИЕ"):
            break
    return count


def _correct_margins(document: Document) -> int:
    section = document.sections[0]
    left_ok = abs(section.left_margin - Mm(30)) <= Mm(3)
    right_ok = abs(section.right_margin - Mm(15)) <= Mm(3)
    top_ok = abs(section.top_margin - Mm(20)) <= Mm(3)
    bottom_ok = abs(section.bottom_margin - Mm(20)) <= Mm(3)
    return int(left_ok and right_ok and top_ok and bottom_ok)


def extract_features(docx_path: str | Path) -> dict[str, float | int]:
    document = Document(str(docx_path))
    texts = _paragraph_texts(document)
    full_text = "\n".join(texts)
    upper_text = full_text.upper()
    lower_text = full_text.lower()
    para_lengths = [len(text) for text in texts]

    keyword_hits = {
        dtype: sum(1 for word in words if word in lower_text)
        for dtype, words in CLASS_KEYWORDS.items()
    }
    best_keyword_score = max(keyword_hits.values(), default=0)

    features = {
        "page_count": max(1, len(document.sections), len(full_text) // 1800),
        "heading_depth": max((_heading_depth(text) for text in texts), default=0),
        "has_toc": int(any("СОДЕРЖАНИЕ" == text.upper() for text in texts)),
        "has_title": int(any(len(text) < 80 for text in texts[:10]) and len(texts[:10]) >= 3),
        "has_abstract": int("РЕФЕРАТ" in upper_text),
        "has_intro": int("ВВЕДЕНИЕ" in upper_text),
        "has_conclusion": int("ЗАКЛЮЧЕНИЕ" in upper_text),
        "table_count": len(document.tables),
        "figure_count": len(re.findall(r"\bРисунок\s+\d+", full_text, flags=re.IGNORECASE)),
        "formula_count": len(re.findall(r"\(\d+(?:\.\d+)*\)\s*$", full_text, flags=re.MULTILINE)),
        "ref_count": _ref_count(texts),
        "appendix_count": len(re.findall(r"\bПРИЛОЖЕНИЕ\b", upper_text)),
        "correct_margins": _correct_margins(document),
        "char_count": len(full_text),
        "avg_para_len": round(sum(para_lengths) / len(para_lengths), 2) if para_lengths else 0.0,
        "keyword_score": float(best_keyword_score),
    }
    return features


def guess_doc_type(features: dict) -> str:
    if features["has_abstract"] and features["page_count"] >= 35:
        return "scientific_rpt"
    if features["appendix_count"] or features["page_count"] >= 45:
        return "thesis_bachelor"
    if features["has_toc"] and features["page_count"] >= 12:
        return "coursework"
    if features["table_count"] >= 2 and features["page_count"] < 12:
        return "lab_work"
    return "practice"

