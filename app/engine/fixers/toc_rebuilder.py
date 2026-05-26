from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def fix(document) -> list[dict]:
    for paragraph in document.paragraphs[:10]:
        if paragraph.text.strip().upper() == "СОДЕРЖАНИЕ":
            return []

    first = document.paragraphs[0]
    toc = first.insert_paragraph_before("")
    title = toc.insert_paragraph_before("СОДЕРЖАНИЕ")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
    _add_toc_field(toc)
    return [{
        "rule_id": "toc_rebuild",
        "paragraph_index": 0,
        "before": "toc missing",
        "after": "inserted Word TOC field placeholder",
    }]
