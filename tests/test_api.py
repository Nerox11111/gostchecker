from pathlib import Path
import base64
import hashlib
import hmac
import os
from urllib.parse import urlencode

os.environ["VK_SECRET_KEY"] = "test_secret"
os.environ["VK_APP_ID"] = "123456"
os.environ.setdefault("DATA_DIR", "./data")
os.environ.setdefault("DB_PATH", "./data/test.sqlite3")
os.environ.setdefault("MODEL_PATH", "./data/models/classifier.joblib")

from docx import Document
from fastapi.testclient import TestClient

from app.main import app


def _vk_launch_params() -> str:
    params = {
        "vk_app_id": "123456",
        "vk_user_id": "777",
        "vk_language": "ru",
        "vk_platform": "desktop_web",
    }
    ordered = urlencode(sorted(params.items()))
    digest = hmac.new(b"test_secret", ordered.encode("utf-8"), hashlib.sha256).digest()
    sign = base64.urlsafe_b64encode(digest).decode("utf-8").rstrip("=")
    return f"?{ordered}&sign={sign}"


def _docx(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("УНИВЕРСИТЕТ")
    doc.add_paragraph("КУРСОВАЯ РАБОТА")
    doc.add_paragraph("СОДЕРЖАНИЕ")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("1 Первый раздел")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("СПИСОК ИСТОЧНИКОВ")
    doc.save(path)
    return path


def test_health_and_check(tmp_path: Path):
    vk_headers = {"X-VK-Sign": _vk_launch_params()}

    with TestClient(app) as client:
        health = client.get("/api/health")
        assert health.status_code == 200
        assert health.json()["status"] == "ok"

        path = _docx(tmp_path / "sample.docx")
        with path.open("rb") as f:
            response = client.post(
                "/api/check",
                headers=vk_headers,
                files={"file": ("sample.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"mode": "medium"},
            )

        assert response.status_code == 200
        payload = response.json()
        assert "session_id" in payload
        assert payload["mode"] == "MEDIUM"

        history = client.get("/api/history", headers=vk_headers)
        assert history.status_code == 200
        assert history.json()["sessions"]
