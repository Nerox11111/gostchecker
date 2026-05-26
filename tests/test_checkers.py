from pathlib import Path

from docx import Document
from docx.shared import Mm, Pt

from app.engine.checkers.font_checker import check_font
from app.engine.checkers.margins_checker import check_margins
from app.engine.checkers.structure_checker import check_required_structure
from app.engine.types import CheckContext
from app.ml.features import extract_features


def _context() -> CheckContext:
    return CheckContext(features={}, mode="MEDIUM")


def test_margins_checker_detects_wrong_margins():
    doc = Document()
    doc.add_paragraph("Тест")
    section = doc.sections[0]
    section.left_margin = Mm(10)

    result = check_margins(doc, _context())

    assert result.score < 1
    assert any(issue.rule_id == "MARGINS" for issue in result.issues)


def test_font_checker_detects_small_font():
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Текст")
    run.font.name = "Arial"
    run.font.size = Pt(10)

    result = check_font(doc, _context())

    assert result.score < 1
    assert {issue.rule_id for issue in result.issues} >= {"FONT_NAME", "FONT_SIZE"}


def test_structure_checker_required_sections():
    doc = Document()
    doc.add_paragraph("СОДЕРЖАНИЕ")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")

    result = check_required_structure(doc, _context())

    assert result.score < 1
    assert any(issue.rule_id == "STRUCTURE" for issue in result.issues)


def test_feature_extraction(tmp_path: Path):
    path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("СОДЕРЖАНИЕ")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("1 Первый раздел")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.save(path)

    features = extract_features(path)

    assert features["has_toc"] == 1
    assert features["has_intro"] == 1
    assert features["heading_depth"] == 1

